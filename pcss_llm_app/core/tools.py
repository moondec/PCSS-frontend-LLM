import os
from typing import Optional
try:
    from langchain_core.tools import tool, StructuredTool
except ImportError:
    from langchain.tools import tool, StructuredTool
from pydantic import BaseModel, Field
from docx import Document
from pypdf import PdfReader
from openai import OpenAI
import base64
import mimetypes
try:
    import pypandoc
except ImportError:
    pypandoc = None

try:
    from ddgs import DDGS
except ImportError:
    from duckduckgo_search import DDGS


# Pydantic schemas for tools with multiple parameters
class SaveDocumentSchema(BaseModel):
    file_path: str = Field(description="Target file name with extension (e.g., 'report.pdf', 'summary.docx')")
    content: str = Field(description="HTML-formatted content (use <h1>, <p>, <ul>, <li>, <b>, <i> tags for formatting)")
    title: str = Field(default="Document", description="Document title (used in HTML header)")

class DocumentTools:
    def __init__(self, root_dir: str):
        self.root_dir = root_dir

    def _get_full_path(self, file_path: str) -> str:
        """Resolve path relative to root_dir and ensure safety."""
        # Simple path join, in production needs traversal protection
        return os.path.join(self.root_dir, file_path)

    # @tool("write_docx")
    def write_docx(self, file_path: str, text: str):
        """
        Creates a new Word document (.docx) with the given text.
        Args:
            file_path: The name of the file to save (e.g., 'document.docx').
            text: The text content to write into the document.
        """
        try:
            full_path = self._get_full_path(file_path)
            doc = Document()
            for line in text.split('\n'):
                doc.add_paragraph(line)
            doc.save(full_path)
            return f"Successfully created DOCX file: {file_path}"
        except Exception as e:
            return f"Error writing DOCX file: {str(e)}"

    # @tool("read_docx")
    def read_docx(self, file_path: str) -> str:
        """
        Reads text content from a Word document (.docx).
        Args:
            file_path: The name of the file to read.
        """
        try:
            full_path = self._get_full_path(file_path)
            if not os.path.exists(full_path):
                return f"Error: File {file_path} not found."
            
            doc = Document(full_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            return '\n'.join(full_text)
        except Exception as e:
            return f"Error reading DOCX file: {str(e)}"

    # @tool("read_pdf")
    def read_pdf(self, file_path: str) -> str:
        """
        Reads text content from a PDF file (.pdf).
        Args:
            file_path: The name of the file to read.
        """
        try:
            full_path = self._get_full_path(file_path)
            if not os.path.exists(full_path):
                return f"Error: File {file_path} not found."
            
            reader = PdfReader(full_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as e:
            return f"Error reading PDF file: {str(e)}"

    def save_document(self, file_path: str, content: str, title: str = "Document") -> str:
        """
        Saves formatted content to a document file. Supports: .pdf, .docx, .html, .txt
        For PDF/DOCX: automatically creates HTML first, then converts via Pandoc.
        Args:
            file_path: Target file name with extension (e.g., 'report.pdf', 'summary.docx').
            content: HTML-formatted content (use <h1>, <p>, <ul>, <li>, <b>, <i> tags for formatting).
            title: Document title (used in HTML header).
        """
        import subprocess
        import tempfile
        
        ext = os.path.splitext(file_path)[1].lower()
        full_path = self._get_full_path(file_path)
        
        # Wrap content in proper HTML structure if not already
        if not content.strip().startswith('<!DOCTYPE') and not content.strip().startswith('<html'):
            html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        h1 {{ color: #333; }}
        h2 {{ color: #555; }}
        ul, ol {{ margin-left: 20px; }}
        li {{ margin-bottom: 10px; }}
        .source {{ font-style: italic; color: #666; }}
    </style>
</head>
<body>
{content}
</body>
</html>"""
        else:
            html_content = content
        
        try:
            # For TXT: just write plain text (strip HTML)
            if ext == '.txt':
                from html import unescape
                import re
                plain_text = re.sub(r'<[^>]+>', '', html_content)
                plain_text = unescape(plain_text)
                with open(full_path, 'w', encoding='utf-8') as f:
                    f.write(plain_text)
                return f"Successfully saved text file: {file_path}"
            
            # For HTML: write directly
            if ext == '.html':
                with open(full_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                return f"Successfully saved HTML file: {file_path}"
            
            # For PDF and DOCX: create temp HTML, then convert
            if ext in ['.pdf', '.docx']:
                # Create temp HTML file
                html_path = full_path.replace(ext, '.html')
                with open(html_path, 'w', encoding='utf-8') as f:
                    f.write(html_content)
                
                # Try conversion
                if ext == '.pdf':
                    # Method 1: weasyprint (best for HTML->PDF, pure Python)
                    try:
                        from weasyprint import HTML
                        HTML(filename=html_path).write_pdf(full_path)
                        if os.path.exists(full_path) and os.path.getsize(full_path) > 0:
                            return f"Successfully saved PDF: {file_path}"
                    except ImportError:
                        pass  # weasyprint not installed
                    except Exception as e:
                        print(f"weasyprint failed: {e}")
                    
                    # Method 2: wkhtmltopdf (fallback)
                    try:
                        result = subprocess.run(
                            ['wkhtmltopdf', '--encoding', 'utf-8', '--quiet', html_path, full_path],
                            capture_output=True, text=True, timeout=60
                        )
                        if result.returncode == 0 and os.path.exists(full_path):
                            return f"Successfully saved PDF: {file_path}"
                    except FileNotFoundError:
                        pass  # wkhtmltopdf not installed
                    except Exception:
                        pass
                    
                    # Method 3: Pandoc with wkhtmltopdf engine
                    if pypandoc:
                        try:
                            pypandoc.convert_file(html_path, 'pdf', outputfile=full_path, 
                                                  extra_args=['--pdf-engine=wkhtmltopdf'])
                            if os.path.exists(full_path) and os.path.getsize(full_path) > 0:
                                return f"Successfully saved PDF: {file_path}"
                        except Exception:
                            pass
                    
                    # Fallback: Create DOCX instead and inform user
                    docx_path = full_path.replace('.pdf', '.docx')
                    if pypandoc:
                        try:
                            pypandoc.convert_file(html_path, 'docx', outputfile=docx_path)
                            return f"PDF conversion failed (missing wkhtmltopdf). Created DOCX instead: {file_path.replace('.pdf', '.docx')}. Install wkhtmltopdf: 'brew install wkhtmltopdf'"
                        except Exception as e:
                            return f"Error: Could not create PDF or DOCX. Details: {str(e)}"
                    return "Error: No PDF conversion tool available. Install wkhtmltopdf or pypandoc."
                
                elif ext == '.docx':
                    if pypandoc:
                        try:
                            pypandoc.convert_file(html_path, 'docx', outputfile=full_path)
                            if os.path.exists(full_path):
                                return f"Successfully saved DOCX: {file_path}"
                        except Exception as e:
                            return f"Error converting to DOCX: {str(e)}"
                    else:
                        # Fallback: use python-docx directly
                        try:
                            from bs4 import BeautifulSoup
                            from docx import Document as DocxDocument
                            
                            soup = BeautifulSoup(html_content, 'html.parser')
                            doc = DocxDocument()
                            doc.add_heading(title, 0)
                            
                            for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'li']):
                                text = element.get_text(strip=True)
                                if element.name.startswith('h'):
                                    level = int(element.name[1])
                                    doc.add_heading(text, level)
                                else:
                                    doc.add_paragraph(text)
                            
                            doc.save(full_path)
                            return f"Successfully saved DOCX: {file_path}"
                        except Exception as e:
                            return f"Error creating DOCX: {str(e)}"
            
            return f"Unsupported format: {ext}. Use .pdf, .docx, .html, or .txt"
            
        except Exception as e:
            return f"Error saving document: {str(e)}"

    def get_tools(self):
        """Returns the list of bound tools."""
        # We need to bind self to the tool methods manually or re-create them?
        # The @tool decorator works on functions. If used on methods, self needs handling.
        # Better approach: Define these as unbound functions or use StructuredTool.from_function
        # wrapping the instance methods.
        
        # Let's use closure-based definition inside get_tools to capture self
        # Or better, just use the methods since I used @tool decorator?
        # Actually @tool on methods might be tricky with 'self' in signature for LangChain.
        
        # Simpler Pattern:
        from langchain_core.tools import StructuredTool

        return [
            StructuredTool.from_function(
                func=self.save_document,
                name="save_document",
                description="Saves formatted content to a document file (.pdf, .docx, .html, .txt). Use HTML tags for formatting (h1, h2, p, ul, li, b, i). For PDF/DOCX: automatically handles conversion. This is the PREFERRED tool for creating formatted documents.",
                args_schema=SaveDocumentSchema
            ),
            StructuredTool.from_function(
                func=self.write_docx,
                name="write_docx",
                description="Creates a new Word document (.docx) with plain text. Use save_document instead for formatted content."
            ),
            StructuredTool.from_function(
                func=self.read_docx,
                name="read_docx",
                description="Reads text content from a Word document (.docx)."
            ),
            StructuredTool.from_function(
                func=self.read_pdf,
                name="read_pdf",
                description="Reads text content from a PDF file (.pdf)."
            ),
        ]

class OCRTools:
    def __init__(self, root_dir: str, api_key: str):
        self.root_dir = root_dir
        self.api_key = api_key
        self.client = OpenAI(
            api_key=api_key,
            base_url="https://llm.hpc.pcss.pl/v1"
        )
        self.model = "Nanonets-OCR-s"

    def _get_full_path(self, file_path: str) -> str:
        return os.path.join(self.root_dir, file_path)

    def ocr_image(self, file_path: str) -> str:
        """
        Extracts text from an image file (PNG, JPG, JPEG) using OCR.
        Args:
            file_path: The path to the image file.
        """
        try:
            full_path = self._get_full_path(file_path)
            if not os.path.exists(full_path):
                return f"Error: File {file_path} not found."
            
            # Basic mime check
            mime_type, _ = mimetypes.guess_type(full_path)
            if not mime_type or not mime_type.startswith('image'):
                return f"Error: File {file_path} does not appear to be an image ({mime_type})."

            with open(full_path, "rb") as image_file:
                base64_image = base64.b64encode(image_file.read()).decode('utf-8')
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": "Transcribe the text in this image verbatim. Output ONLY the text."},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error performing OCR: {str(e)}"

    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.ocr_image,
                name="ocr_image",
                description="Extracts text from an image file (PNG, JPG) using OCR. Use this to read scanned documents or text in images."
            )
        ]

class PandocTools:
    def __init__(self, root_dir: str):
        self.root_dir = root_dir

    def _get_full_path(self, file_path: str) -> str:
         return os.path.join(self.root_dir, file_path)

    def convert_document(self, source_path: str, output_format: str) -> str:
        """
        Converts a document (e.g. HTML) to another format (e.g. DOCX, PDF) using Pandoc.
        Args:
            source_path: Path to the source file (e.g. 'report.html').
            output_format: Target format extension (e.g. 'docx', 'pdf').
        """
        try:
            full_source = self._get_full_path(source_path)
            if not os.path.exists(full_source):
                 return f"Error: Source file {source_path} not found."
            
            # Construct output filename
            base_name = os.path.splitext(source_path)[0]
            target_filename = f"{base_name}.{output_format}"
            full_target = self._get_full_path(target_filename)
            
            # For HTML to PDF, try wkhtmltopdf first (more reliable than LaTeX)
            if output_format.lower() == 'pdf' and source_path.lower().endswith('.html'):
                try:
                    import subprocess
                    result = subprocess.run(
                        ['wkhtmltopdf', '--encoding', 'utf-8', full_source, full_target],
                        capture_output=True, text=True, timeout=30
                    )
                    if result.returncode == 0:
                        return f"Successfully converted {source_path} to {target_filename}."
                except FileNotFoundError:
                    pass  # wkhtmltopdf not installed, fall through to pypandoc
                except Exception:
                    pass
            
            # Fallback to pypandoc
            if pypandoc is None:
                return "Error: pypandoc module is not installed and wkhtmltopdf not available."
            
            output = pypandoc.convert_file(full_source, output_format, outputfile=full_target)
            return f"Successfully converted {source_path} to {target_filename}."
        except Exception as e:
             return f"Error converting document: {str(e)}"

    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.convert_document,
                name="convert_document",
                description="Converts documents between formats (e.g. HTML to DOCX). Best used for creating formatted reports: 'Write content to .html then convert to .docx'."
            )
        ]


class VisionTools:
    def __init__(self, root_dir: str, api_key: str):
        self.root_dir = root_dir
        self.api_key = api_key
        # Always use GPT-4o for vision tasks, regardless of main agent model
        self.client = OpenAI(
            api_key=api_key,
            base_url="https://llm.hpc.pcss.pl/v1"
        )
        self.model = "gpt-4o"

    def _get_full_path(self, file_path: str) -> str:
        return os.path.join(self.root_dir, file_path)

    def analyze_image(self, file_path: str, prompt: str = "Describe this image in detail.") -> str:
        """
        Analyzes an image file using a Vision LLM (GPT-4o) to understand its content, layout, or extract data.
        Args:
            file_path: The name of the image file (e.g., 'chart.png').
            prompt: Question or instruction about the image (e.g., 'What is the trend in this chart?').
        """
        try:
            full_path = self._get_full_path(file_path)
            if not os.path.exists(full_path):
                return f"Error: File {file_path} not found."
            
            # Basic mime check
            mime_type, _ = mimetypes.guess_type(full_path)
            if not mime_type or not mime_type.startswith('image'):
                return f"Error: File {file_path} does not appear to be an image ({mime_type})."

            with open(full_path, "rb") as image_file:
                base64_image = base64.b64encode(image_file.read()).decode('utf-8')
            
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:{mime_type};base64,{base64_image}"
                                }
                            }
                        ]
                    }
                ]
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"Error analyzing image: {str(e)}"

    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.analyze_image,
                name="analyze_image",
                description="Analyzes an image using GPT-4o. Use this to descriptive scenes, understand charts, or analyze document layouts. Input: file_path and prompt."
            )
        ]


class ChartTools:
    """
    Tools for generating charts and visualizations using matplotlib.
    """
    def __init__(self, root_dir: str = "."):
        self.root_dir = root_dir
    
    def generate_chart(
        self, 
        chart_type: str, 
        data: str, 
        labels: str, 
        file_path: str,
        title: str = "Chart",
        x_label: str = "",
        y_label: str = "",
        colors: str = ""
    ) -> str:
        """
        Generate a chart and save it as PNG/JPG.
        
        Args:
            chart_type: Type of chart - 'bar', 'line', 'pie', 'scatter', 'horizontal_bar'
            data: Comma-separated values (e.g., "10,25,30,15,20")
            labels: Comma-separated labels (e.g., "Jan,Feb,Mar,Apr,May")
            file_path: Output file path (e.g., "charts/sales.png")
            title: Chart title
            x_label: X-axis label (for bar/line charts)
            y_label: Y-axis label (for bar/line charts)
            colors: Optional comma-separated colors (e.g., "#FF6384,#36A2EB,#FFCE56")
        
        Returns:
            Success message with file path or error message.
        """
        try:
            import matplotlib
            matplotlib.use('Agg')  # Non-interactive backend
            import matplotlib.pyplot as plt
            import numpy as np
            
            # Parse data
            try:
                values = [float(x.strip()) for x in data.split(",")]
            except ValueError:
                return "Error: 'data' must be comma-separated numbers (e.g., '10,25,30')"
            
            # Parse labels
            label_list = [x.strip() for x in labels.split(",")]
            if len(label_list) != len(values):
                return f"Error: Number of labels ({len(label_list)}) must match number of data points ({len(values)})"
            
            # Parse colors
            if colors:
                color_list = [x.strip() for x in colors.split(",")]
            else:
                # Default color palette
                color_list = plt.cm.Set3.colors[:len(values)]
            
            # Create figure
            fig, ax = plt.subplots(figsize=(10, 6))
            
            chart_type_lower = chart_type.lower().strip()
            
            if chart_type_lower == "bar":
                bars = ax.bar(label_list, values, color=color_list[:len(values)])
                # Add value labels on bars
                for bar, val in zip(bars, values):
                    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 0.5, 
                           f'{val:.1f}', ha='center', va='bottom', fontsize=9)
                           
            elif chart_type_lower == "horizontal_bar":
                bars = ax.barh(label_list, values, color=color_list[:len(values)])
                for bar, val in zip(bars, values):
                    ax.text(bar.get_width() + 0.5, bar.get_y() + bar.get_height()/2,
                           f'{val:.1f}', ha='left', va='center', fontsize=9)
                           
            elif chart_type_lower == "line":
                ax.plot(label_list, values, marker='o', linewidth=2, markersize=8, 
                       color=color_list[0] if colors else '#36A2EB')
                ax.fill_between(label_list, values, alpha=0.3)
                
            elif chart_type_lower == "pie":
                ax.pie(values, labels=label_list, autopct='%1.1f%%', 
                      colors=color_list[:len(values)], startangle=90)
                ax.axis('equal')
                
            elif chart_type_lower == "scatter":
                # For scatter, use index as X if no separate X data
                x_vals = range(len(values))
                ax.scatter(x_vals, values, c=color_list[:len(values)], s=100)
                ax.set_xticks(list(x_vals))
                ax.set_xticklabels(label_list)
                
            else:
                return f"Error: Unknown chart_type '{chart_type}'. Use: bar, line, pie, scatter, horizontal_bar"
            
            # Set labels and title
            ax.set_title(title, fontsize=14, fontweight='bold', pad=15)
            if x_label and chart_type_lower != "pie":
                ax.set_xlabel(x_label, fontsize=11)
            if y_label and chart_type_lower != "pie":
                ax.set_ylabel(y_label, fontsize=11)
            
            # Style
            if chart_type_lower != "pie":
                ax.spines['top'].set_visible(False)
                ax.spines['right'].set_visible(False)
                ax.grid(axis='y', alpha=0.3)
            
            plt.tight_layout()
            
            # Resolve path
            full_path = os.path.join(self.root_dir, file_path)
            os.makedirs(os.path.dirname(full_path) if os.path.dirname(full_path) else ".", exist_ok=True)
            
            # Determine format from extension
            ext = os.path.splitext(file_path)[1].lower()
            if ext not in ['.png', '.jpg', '.jpeg', '.svg', '.pdf']:
                full_path += '.png'
            
            plt.savefig(full_path, dpi=150, bbox_inches='tight', 
                       facecolor='white', edgecolor='none')
            plt.close(fig)
            
            return f"Chart saved successfully: {file_path}"
            
        except ImportError:
            return "Error: matplotlib is required. Install with: pip install matplotlib"
        except Exception as e:
            return f"Error generating chart: {str(e)}"
    
    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.generate_chart,
                name="generate_chart",
                description="""Generate a chart/visualization and save as PNG/JPG.
Args:
- chart_type: 'bar', 'line', 'pie', 'scatter', or 'horizontal_bar'
- data: Comma-separated values (e.g., "10,25,30,15")
- labels: Comma-separated labels (e.g., "Q1,Q2,Q3,Q4")
- file_path: Output path (e.g., "chart.png")
- title: Chart title
- x_label, y_label: Optional axis labels
- colors: Optional comma-separated hex colors"""
            )
        ]


class WebSearchTools:
    """
    Optimized web search tools for deep research.
    """
    def __init__(self, api_key: str = None, model_name: str = "gpt-4o", base_url: str = "https://llm.hpc.pcss.pl/v1"):
        self.api_key = api_key
        self.model_name = model_name
        self.base_url = base_url
        self.client = None
        if self.api_key:
            try:
                self.client = OpenAI(api_key=self.api_key, base_url=self.base_url)
            except Exception:
                pass

    def search_web(self, query: str, max_results: int = 10) -> str:
        """
        Performs a general web search using DuckDuckGo.
        Best for: definitions, how-to guides, reference information.
        Args:
            query: The search query string.
            max_results: Maximum number of results (default 10).
        """
        try:
            results = []
            with DDGS() as ddgs:
                ddgs_gen = ddgs.text(query, region="pl-pl", max_results=max_results)
                if ddgs_gen:
                    for i, r in enumerate(ddgs_gen, 1):
                        results.append(
                            f"[{i}] {r.get('title', 'No title')}\n"
                            f"    URL: {r.get('href', '')}\n"
                            f"    {r.get('body', '')[:200]}"
                        )
            
            if not results:
                return "No results found. Try rephrasing your query."
            
            return "SEARCH RESULTS:\n\n" + "\n\n".join(results) + "\n\n[TIP: Use visit_page on URLs to read full content]"
        except Exception as e:
            return f"Search error: {str(e)}"

    def search_news(self, query: str, max_results: int = 8) -> str:
        """
        Searches for recent NEWS articles using DuckDuckGo News.
        Best for: current events, breaking news, recent developments.
        Args:
            query: The news search query.
            max_results: Maximum number of news articles (default 8).
        """
        try:
            results = []
            with DDGS() as ddgs:
                news_gen = ddgs.news(query, region="pl-pl", max_results=max_results)
                if news_gen:
                    for i, r in enumerate(news_gen, 1):
                        date = r.get('date', 'Unknown date')
                        source = r.get('source', 'Unknown source')
                        results.append(
                            f"[{i}] {r.get('title', 'No title')}\n"
                            f"    Source: {source} | Date: {date}\n"
                            f"    URL: {r.get('url', '')}\n"
                            f"    {r.get('body', '')[:200]}"
                        )
            
            if not results:
                return "No news found. Try broader search terms."
            
            return "NEWS RESULTS:\n\n" + "\n\n".join(results) + "\n\n[TIP: Use visit_page on URLs to read full articles]"
        except Exception as e:
            return f"News search error: {str(e)}"

    def visit_page(self, url: str) -> str:
        """
        Visits a URL and extracts the main article content.
        Uses readability algorithm to extract relevant text, ignoring navigation/ads.
        Args:
            url: The URL to visit.
        """
        try:
            import requests
            from readability import Document
            
            headers = {
                "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
                "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                "Accept-Language": "pl-PL,pl;q=0.9,en;q=0.8",
            }
            
            response = requests.get(url, headers=headers, timeout=15)
            response.raise_for_status()
            response.encoding = response.apparent_encoding or 'utf-8'
            
            # Use readability to extract main content
            doc = Document(response.text)
            title = doc.title()
            
            # Get clean HTML content and convert to text
            from bs4 import BeautifulSoup
            content_html = doc.summary()
            soup = BeautifulSoup(content_html, 'html.parser')
            
            # Extract text with better formatting
            text_parts = []
            for elem in soup.find_all(['p', 'h1', 'h2', 'h3', 'li']):
                text = elem.get_text(strip=True)
                if text and len(text) > 20:  # Skip short fragments
                    text_parts.append(text)
            
            content = "\n\n".join(text_parts)
            
            if not content or len(content) < 100:
                # Fallback to basic extraction if readability fails
                soup = BeautifulSoup(response.text, 'html.parser')
                for tag in soup(['script', 'style', 'nav', 'footer', 'header', 'aside']):
                    tag.decompose()
                content = soup.get_text(separator='\n', strip=True)
            
            # Limit to ~10000 chars (roughly 2500 words)
            if len(content) > 10000:
                content = content[:10000] + "\n\n[... Content truncated ...]"
            
            return f"=== {title} ===\nSource: {url}\n\n{content}"
            
        except ImportError as ie:
            return f"Missing library: {ie}. Install with: pip install readability-lxml requests beautifulsoup4"
        except requests.exceptions.Timeout:
            return f"Timeout: Page took too long to load: {url}"
        except requests.exceptions.HTTPError as he:
            return f"HTTP Error {he.response.status_code}: Cannot access {url}"
        except Exception as e:
            return f"Error visiting page: {str(e)}"

    def get_tools(self):
        return [
            StructuredTool.from_function(
                func=self.search_web,
                name="search_web",
                description="Search the web for general information (definitions, guides, reference). Returns links with snippets. Use visit_page to read full content."
            ),
            StructuredTool.from_function(
                func=self.search_news,
                name="search_news",
                description="Search for recent NEWS and current events. Returns news articles with dates and sources. Use visit_page to read full articles."
            ),
            StructuredTool.from_function(
                func=self.visit_page,
                name="visit_page",
                description="Visit a URL and extract the main article text. Use AFTER search_web or search_news to read full content from a link."
            )
        ]

